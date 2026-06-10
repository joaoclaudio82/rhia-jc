"""Gera currículos de teste em PDF (nato-digital) e DOCX (com tabela)."""
from pathlib import Path

AQUI = Path(__file__).parent
TEXTO = (AQUI / "cv_exemplo.txt").read_text(encoding="utf-8")


def gerar_pdf() -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=11)
    for linha in TEXTO.splitlines():
        linha = linha.replace("\u2014", "-")  # fonte core não suporta em-dash
        if linha.strip():
            pdf.multi_cell(0, 6, linha, new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.ln(4)
    pdf.output(str(AQUI / "cv_exemplo.pdf"))


def gerar_docx() -> None:
    import docx

    doc = docx.Document()
    doc.add_heading("Carlos Eduardo Lima", level=0)
    doc.add_paragraph("Analista de Dados")
    doc.add_paragraph(
        "Email: carlos.lima@example.com | Telefone: (21) 91234-5678 | Rio de Janeiro - RJ"
    )
    doc.add_heading("Experiência", level=1)

    tabela = doc.add_table(rows=3, cols=4)
    cabecalho = ["Cargo", "Empresa", "Período", "Descrição"]
    for i, txt in enumerate(cabecalho):
        tabela.rows[0].cells[i].text = txt
    tabela.rows[1].cells[0].text = "Analista de Dados Pleno"
    tabela.rows[1].cells[1].text = "FinBank"
    tabela.rows[1].cells[2].text = "2022 - atual"
    tabela.rows[1].cells[3].text = "Dashboards em Power BI e modelagem em SQL."
    tabela.rows[2].cells[0].text = "Analista de Dados Júnior"
    tabela.rows[2].cells[1].text = "VarejoTop"
    tabela.rows[2].cells[2].text = "2019 - 2022"
    tabela.rows[2].cells[3].text = "ETL com Python e pandas."

    doc.add_heading("Formação", level=1)
    doc.add_paragraph("Estatística — UFRJ (2015 - 2019) — concluído")
    doc.add_heading("Idiomas", level=1)
    doc.add_paragraph("Português (nativo), Inglês (intermediário)")
    doc.save(str(AQUI / "cv_exemplo.docx"))


if __name__ == "__main__":
    gerar_pdf()
    gerar_docx()
    print("Fixtures geradas:", AQUI / "cv_exemplo.pdf", AQUI / "cv_exemplo.docx")
